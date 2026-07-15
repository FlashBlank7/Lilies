from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/lilies_agent_scenario_capability_boundary_v0_4_x_revised.docx"
OUTPUT = ROOT / "docs/lilies_agent_scenario_capability_boundary_v0_4_x_latest.docx"

ANNOTATION_COVERAGE = [
    {
        "id": "ANN-001",
        "source_contains": "harness本身工作流里也有",
        "output_markers": ["Evaluation Harness Profile", "Quick、Guided、Governed"],
    },
    {
        "id": "ANN-002",
        "source_contains": "客户点开积木发现一堆json",
        "output_markers": ["面向人的积木配置界面", "发布权属于用户", "一键把失败原因"],
    },
    {
        "id": "ANN-003",
        "source_contains": "多种heavy程度的框架需要并存",
        "output_markers": ["三种并存的构建、验收与监管模式", "独立 Governance Console"],
    },
    {
        "id": "ANN-004",
        "source_contains": "以上才是真正的原始设定",
        "output_markers": ["Lilies 的原始设定是", "Program Charter"],
    },
    {
        "id": "ANN-005",
        "source_contains": "模版市场的设想",
        "output_markers": ["模板市场与工作流模块", "Template/Module Registry"],
    },
    {
        "id": "ANN-006",
        "source_contains": "是不是所有的缺的都需要体现在积木里",
        "output_markers": ["先判断承载层，再决定是否新增积木", "五类承载位置"],
    },
    {
        "id": "ANN-007",
        "source_contains": "对批注问题问答的回答：没问题",
        "output_markers": ["对层次模型问题的直接回答  没问题"],
    },
]


def document_fragments(document: Document) -> list[str]:
    fragments = [paragraph.text for paragraph in document.paragraphs]
    fragments.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return fragments


def inline_annotation_inventory(document: Document) -> list[str]:
    annotations: list[str] = []
    for fragment in document_fragments(document):
        if fragment.startswith("批注："):
            annotations.append(fragment)
        elif "\n批注：" in fragment:
            annotations.append("批注：" + fragment.split("\n批注：", 1)[1])
        if fragment.startswith("对批注问题问答的回答："):
            annotations.append(fragment)
    return annotations


def validate_annotation_inventory(document: Document) -> list[str]:
    annotations = inline_annotation_inventory(document)
    matches: dict[str, str] = {}
    for item in ANNOTATION_COVERAGE:
        found = [annotation for annotation in annotations if item["source_contains"] in annotation]
        if len(found) != 1:
            raise ValueError(f"{item['id']} expected one source annotation, found {len(found)}")
        matches[item["id"]] = found[0]
    if len(matches) != len(annotations):
        raise ValueError(
            f"annotation inventory has {len(annotations)} entries but coverage maps {len(matches)}"
        )
    return annotations


def validate_annotation_output(document: Document) -> None:
    fragments = document_fragments(document)
    output_text = "\n".join(fragments)
    residue = [
        fragment
        for fragment in fragments
        if "批注：" in fragment or fragment.startswith("对批注问题问答的回答：")
    ]
    if residue:
        raise ValueError(f"inline annotation residue remains: {residue[0][:80]}")
    for item in ANNOTATION_COVERAGE:
        for marker in item["output_markers"]:
            if marker not in output_text:
                raise ValueError(f"{item['id']} output marker missing: {marker}")


def paragraph_starting(document: Document, prefix: str):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise ValueError(f"paragraph not found: {prefix}")


def replace_paragraph(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    paragraph.add_run(text)


def remove_paragraph(paragraph) -> None:
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def insert_paragraph_before(anchor, text: str, style: str = "Normal"):
    element = OxmlElement("w:p")
    anchor._p.addprevious(element)
    paragraph = anchor._parent.add_paragraph()
    paragraph._p.getparent().remove(paragraph._p)
    element.addnext(paragraph._p)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def insert_paragraph_after(anchor, text: str, style: str = "Normal"):
    paragraph = anchor._parent.add_paragraph()
    paragraph._p.getparent().remove(paragraph._p)
    anchor._p.addnext(paragraph._p)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def insert_table_before(document: Document, anchor, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
                if row_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[index] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            if row_index == 0:
                shd = tc_pr.find(qn("w:shd"))
                if shd is None:
                    shd = OxmlElement("w:shd")
                    tc_pr.append(shd)
                shd.set(qn("w:fill"), "F2F4F7")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    anchor._p.addprevious(table._tbl)
    return table


def replace_cell_text(cell, text: str) -> None:
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.space_before = Pt(3)


def table_with_headers(document: Document, expected: list[str]):
    for table in document.tables:
        if not table.rows:
            continue
        actual = [cell.text.strip() for cell in table.rows[0].cells]
        if actual == expected:
            return table
    raise ValueError(f"table not found: {expected}")


def append_source_row(table, source: str, purpose: str, url: str) -> None:
    cells = table.add_row().cells
    values = [source, purpose, url]
    for index, value in enumerate(values):
        replace_cell_text(cells[index], value)


def main() -> None:
    document = Document(SOURCE)
    validate_annotation_inventory(document)

    replace_paragraph(
        document.paragraphs[2],
        "批注吸收与落地控制版：能力组合、分级构建、可选验收与可恢复演进",
    )
    replace_paragraph(document.paragraphs[3], "版本：v0.4.x 架构研究与实施基线（latest）")
    replace_paragraph(document.paragraphs[4], "日期：2026-07-16")
    replace_paragraph(
        document.paragraphs[6],
        "依据：用户正文批注、当前 repo 代码/阶段报告、定向回归证据、外部一手技术资料与 OpenAI 长任务官方建议",
    )
    replace_paragraph(
        document.paragraphs[7],
        "边界：本报告既是能力边界分析，也是后续实现的意图基线；实际完成状态必须由稳定意图 ID、代码、测试、真实运行和阶段闭环共同证明。",
    )
    replace_paragraph(
        document.paragraphs[8],
        "本次重构的核心纠偏  Lilies 的产品原点是模型智能优先：Builder Team 阅读积木手册、结合先验知识和客户需求，用功能可靠的积木、可复用模块与必要的平台服务搭建工作流；Harness 既可以作为工作流内可组合积木，也可以存在于构建/验收过程和平台硬边界。工程化能力应增强而不是阻断交付，轻量、引导和治理模式必须并存。",
    )
    replace_paragraph(
        document.paragraphs[10],
        "前两版报告的根本问题不仅是建模对象错位，还忽略了 Lilies 的原始产品设定和自动演进过程中已经发生的意图偏移。目标工作流、Builder Team、工作流内 Harness、验收 Harness、Runtime 与 Platform Harness 必须分别建模；同时，客户是否运行验收、是否依据验收结果修改、是否发布，应当是可配置的产品选择，而不是统一强制门禁。",
    )
    replace_paragraph(
        document.paragraphs[14],
        "Builder Team 是元系统，不是用户生成的智能体。它应在保留现有直接构建路径的前提下，按能力集合、风险和用户选择提供不同构建与验收模式；自动化新路径必须先以实验并行运行，证明确有改进后再升级为默认。",
    )
    replace_paragraph(
        document.paragraphs[15],
        "画布不应继续承载所有工程控制，但积木也不能只暴露 JSON。业务流程留在画布；每个可调积木提供与其能力匹配的表单、选择器和结构化编辑器；模块市场承载可复用子工作流；预算、队列、秘密、沙盒和审计等不可绕过控制留在平台侧。",
    )
    replace_paragraph(
        document.paragraphs[16],
        "当前 Lilies 已有 Platform Harness API、任务记录、调用次数预算、策略控制和应用内 Monitor 标签，但生成、编辑、运行和验收仍然拥挤。监管能力应按风险分层，简单业务可快速生成和运行，复杂或高风险业务再逐步启用证据、策略与监管；平台能力存在、默认启用和用户可见必须分开陈述。",
    )

    decisions = document.tables[0]
    replace_cell_text(
        decisions.rows[3].cells[2],
        "把构建/验收阶段的 Evaluation Harness 改为可选择的 Profile；保留直接构建和原有流程，新自动化路径先并行实验，验证优于旧路径后再考虑默认化。简单需求允许从需求直接得到可用工作流。",
    )
    replace_cell_text(
        decisions.rows[4].cells[2],
        "画布保留业务步骤、可组合 Harness 与模块引用；可调积木必须提供面向人的配置界面并写回全局 WorkflowSpec。任何影响行为的修改使旧验收证据过期并建议重验，但发布权属于用户。失败用例应产生可执行建议，并可一键交回 Builder Team 通过自然语言工作流编辑修复。",
    )
    replace_cell_text(
        decisions.rows[5].cells[2],
        "建设分层而非一刀切的监管体系：Quick、Guided、Governed 三种模式共用 Builder Team。先解决工作流生成、编辑和运行的可用性，再让重型治理按风险逐步启用；全局监管控制台与 Capability Evidence Registry 服务工程和高风险场景，不阻塞普通客户使用。",
    )

    background = paragraph_starting(document, "Lilies 的初始方向是以较强的模型智能为核心")
    replace_paragraph(
        background,
        "Lilies 的原始设定是：以较强模型智能为核心，Builder Team 阅读积木使用手册并结合先验知识与需求，用功能可靠的积木搭建工作流；同时探索把 Harness 做成积木，使生成的工作流本身更可靠；在搭建和验收过程中加入工程化手段；工作流运行只接受与风险相称的平台监管，当前最直接的边界是避免异常 token 和性能消耗。自动演进过程中产生的大量治理、验收和工程资产有价值，但它们不应反过来绑架工作流的基本生成与使用路径。",
    )

    inline_prefixes = [
        "批注：Lilies 的初始方向",
        "批注：莉莉丝其实有个模版市场",
        "批注：这个部分写的相当不错",
        "对批注问题问答的回答：",
    ]
    for prefix in inline_prefixes:
        remove_paragraph(paragraph_starting(document, prefix))

    requirement_formula = paragraph_starting(document, "C(R) = closure(F)")
    placement_heading = insert_paragraph_after(
        requirement_formula,
        "3.2 能力实现位置：先判断承载层，再决定是否新增积木",
        "Heading 2",
    )
    placement_paragraph = insert_paragraph_after(
        placement_heading,
        "能力映射不能默认等于新增原子积木。每项能力先在五类承载位置中选择：原子积木、可复用模块/子工作流、Runtime 或平台服务、不可绕过的平台控制、外部 Connector/环境契约。选择标准是客户是否需要在业务流程中理解和组合它、是否需要跨工作流复用、是否必须不可删除、是否涉及共享状态，以及是否依赖外部身份和系统。",
    )
    insert_paragraph_after(
        placement_paragraph,
        "模板市场是能力组合层：经过运行和验收的工作流可以封装成带输入输出契约、版本、证据级别和依赖声明的模块，被 Builder Team 作为高阶能力复用。它与原子积木目录并存，避免把所有新能力都拆成越来越多的低层节点。",
    )

    for old, new in [
        ("3.2 能力基元目录", "3.3 能力基元目录"),
        ("3.3 四条关键依赖链", "3.4 四条关键依赖链"),
        ("3.4 集合模型何时合理、何时不够", "3.5 集合模型何时合理、何时不够"),
    ]:
        replace_paragraph(paragraph_starting(document, old), new)

    direct_answer = paragraph_starting(document, "对批注问题的直接回答")
    replace_paragraph(
        direct_answer,
        "对层次模型问题的直接回答  没问题：越来越完整的运行系统能够包揽更多类型的问题，但前提是“更完整”被严格定义为累积的平台运行责任，而不是模糊的复杂度评分。外层包含内层的执行保证，却仍不能自动产生任意领域工具、客户数据契约或高风险许可，因此层次模型必须与能力集合共同使用。",
    )

    replace_paragraph(paragraph_starting(document, "6.1 Harness Profile"), "6.1 Evaluation Harness Profile，而不是固定或强制 Harness")
    profile_intro = paragraph_starting(document, "Builder Team 的智能价值")
    insert_paragraph_after(
        profile_intro,
        "这里的 Profile 专指构建和验收阶段的证据编排，不取代工作流内 Harness 积木，也不等同于 Platform Harness。Profile 是否启用、执行到哪一级、是否自动运行由用户、风险和可用环境共同决定。",
    )

    section_seven = paragraph_starting(document, "7. 架构修正")
    insert_paragraph_before(section_seven, "6.4 三种并存的构建、验收与监管模式", "Heading 2")
    insert_table_before(
        document,
        section_seven,
        ["模式", "适用场景", "默认行为", "用户控制"],
        [
            ["Quick", "简单、低风险、一次性或个人工作流", "需求直接生成可运行草稿；基础结构检查；不强制真实验收", "可立即运行或发布，也可主动升级验收"],
            ["Guided", "需要可靠交付但风险可控", "生成建议用例、可选自动运行、失败建议和一键 Builder 修复", "用户决定修复、重验或带提示发布"],
            ["Governed", "外部副作用、长期任务、多租户或高风险流程", "证据等级、权限、预算、策略和监管按声明范围启用", "高风险边界仍由平台强制，业务发布策略可配置"],
        ],
        [0.9, 2.0, 2.4, 1.2],
    )
    insert_paragraph_before(section_seven, "6.5 验收、修改与发布契约", "Heading 2")
    for text in [
        "验收是决策辅助和可靠性证据，不是所有工作流统一的发布许可证。Quick/Guided 模式允许用户在看见验收状态和风险提示后自行发布；只有明确进入 Governed 模式并绑定组织策略时，平台才执行不可绕过的门禁。",
        "任何影响运行行为的积木配置、连线、输入输出契约或模型参数修改都会使对应 tested_hash/证据标记为 stale，并提示重新验收，但不会删除草稿或隐藏发布操作。",
        "失败用例必须生成结构化修复建议。用户可以一键把失败原因、相关节点、Trace、当前 WorkflowSpec 和建议交回 Builder Team，通过已经存在的自然语言工作流编辑能力生成可预览补丁；应用后再选择重验。",
    ]:
        insert_paragraph_before(section_seven, text)

    governance_heading = paragraph_starting(document, "8. 平台监管控制台")
    insert_paragraph_before(governance_heading, "7.4 面向人的积木配置界面", "Heading 2")
    for text in [
        "工程师点击积木时看到的是按 block schema 生成的领域配置器，而不是原始 JSON。LLM 积木至少提供输入映射、模型、温度、输出 schema 和高级选项；HTTP/工具积木提供方法、地址、认证引用、参数和错误策略；Loop 提供状态、条件、上限、break/cancel 和 checkpoint。",
        "高级 JSON 只作为可折叠的专家模式存在。保存配置必须经过 schema 校验并写回唯一 WorkflowSpec；画布、自然语言说明、客户运行表单和 Builder Team 都读取同一份配置，避免局部 UI 与真实执行分叉。",
    ]:
        insert_paragraph_before(governance_heading, text)
    insert_paragraph_before(governance_heading, "7.5 模板市场与工作流模块", "Heading 2")
    insert_paragraph_before(
        governance_heading,
        "模板市场保存经过验证的 WorkflowSpec 或子流程模块，而不是只保存提示词。模块声明输入、输出、能力依赖、所需运行闭包、配置面板、版本兼容、证据级别和已知边界；Builder Team 优先检索可复用模块，再决定组合现有积木、补充新积木或请求平台能力。",
    )

    roadmap = table_with_headers(document, ["优先级", "工作包", "核心内容", "完成标准"])
    headers = [cell.text for cell in roadmap.rows[0].cells]
    roadmap_rows = [
        ["P0", "Product Intent & Evolution Control", "Program Charter、稳定意图 ID、Stage Contract、偏移登记、独立闭环审计和恢复协议。", "长迭代中任务不丢失、不自行降标；中断后从同一任务恢复。"],
        ["P0", "Quick/Guided/Governed Modes", "保留直接构建；验收自动化可选；按风险启用证据和平台监管。", "简单需求可直接得到工作流，重型治理不再阻塞普通使用。"],
        ["P0", "Human-readable Block Configuration", "schema 驱动配置器、专家 JSON、统一保存、修改后证据 stale、建议重验。", "常用积木无需编辑 JSON；配置真实作用于运行。"],
        ["P0", "Advisory Acceptance & Repair Loop", "发布权回到用户；失败建议和一键交回 Builder 自然语言编辑；Governed 策略例外。", "失败不会只显示红色状态；可预览修复并选择重验。"],
        ["P0", "Capability Requirement Schema", "输出 F/G/X、requires/excludes、required_envelope、external_contracts，并保留原自然语言。", "Codex-like、爬虫和嵌入场景得到不同且可追踪的能力闭包。"],
        ["P0", "Builder Capability Build Contract", "Builder 生成平台覆盖、证据计划、claim_scope 和能力承载位置。", "缺环境时输出 component_verified/blocked，不伪造整体验收。"],
        ["P1", "Template/Module Registry", "Workflow-as-module、输入输出契约、版本、依赖和证据；为模板市场提供后端基础。", "Builder 能复用高阶模块，不必把所有能力都新增为原子积木。"],
        ["P1", "前端三界面边界", "Customer Runtime、Engineer Studio、Governance Console 独立信息架构。", "客户不面对工程噪音；工程师和监管角色各自获得完整信息。"],
        ["P1", "监管控制台 MVP", "任务、策略、worker、queue、token/cost、失败、预算、父子 Trace 与 Evidence Registry。", "平台声明可从 UI 下钻到 API、测试和最近运行。"],
        ["P1", "结构化 Loop 核心化", "显式 state/condition/max/break/cancel/checkpoint/trace；支持工具结果反馈。", "Codex-like 循环不依赖黑箱节点或任意有环连线。"],
        ["P1", "Eval Lab 与 Evaluation Harness Profiles", "mock/contract/sandbox/live/production evidence 协议和可选自动化。", "自动用例与风险、用户选择和可获得证据一致。"],
        ["P1", "Durable Job Substrate", "队列、history、timer、resume、retry、idempotency、worker lease 和告警。", "定时与长期任务可恢复运行。"],
        ["P2", "Connector/Embedding SDK", "身份/租户、schema、webhook、writeback/compensation 和部署 profile。", "客户系统嵌入从项目制原型转为平台能力。"],
        ["P2", "高风险自治治理", "风险分级、预授权、不可绕过确认、紧急停止、域审计和生产演练。", "仅在明确领域和硬控制下开放高风险自治。"],
    ]
    while len(roadmap.rows) > 1:
        roadmap._tbl.remove(roadmap.rows[-1]._tr)
    for row_values in roadmap_rows:
        cells = roadmap.add_row().cells
        for index, value in enumerate(row_values):
            replace_cell_text(cells[index], value)
    for index, value in enumerate(headers):
        replace_cell_text(roadmap.rows[0].cells[index], value)
        for run in roadmap.rows[0].cells[index].paragraphs[0].runs:
            run.bold = True

    conclusion = paragraph_starting(document, "Lilies 不应被定义成")
    replace_paragraph(
        conclusion,
        "Lilies 不应被定义成“能画很多积木的工作流工具”，也不能变成“先通过一套沉重工程门禁才能使用”的平台。它的产品内核仍然是模型智能理解需求并组合可靠能力：原子积木、可复用模块、Runtime 服务、平台硬控制和外部 Connector 各归其位。Quick 模式保护低摩擦生成与运行，Guided 模式提供可选证据和一键修复，Governed 模式承担真正需要的平台监管。",
    )
    final_route = paragraph_starting(document, "一句话路线")
    replace_paragraph(
        final_route,
        "一句话路线  恢复模型智能优先和可靠积木组合的产品原点，用能力闭包决定需要什么，用承载位置决定它应是积木、模块还是平台能力，用分级模式决定需要多少验收和监管，再用可恢复、可审计的自动演进协议把长期设想逐项落地。",
    )

    version_method = paragraph_starting(document, "不要把上述路线压进一个版本")
    insert_paragraph_after(version_method, "10.2 自动演进的长任务控制架构", "Heading 2")
    auto_heading = paragraph_starting(document, "10.2 自动演进的长任务控制架构")
    anchor = auto_heading
    for text in [
        "报告与批注是需求证据，Program Charter 是不可随阶段改写的意图约束，stage-report 仍是唯一下一阶段任务来源。每个报告结论使用稳定 INTENT/CAP/UX/GOV/EVOL ID，并在每份 stage-report 中携带覆盖状态，防止递归摘要让任务消失。",
        "每个版本开始前冻结 Stage Contract：目标、mandatory 任务、验收标准、证据和允许偏移。技术路线可以调整，但删除 mandatory 任务、降低验收、改变目标客户或产品边界必须请求用户决定。blocked 不等于 completed，也不能单独支持版本晋级。",
        "实施者不能以自己的总结作为完成证明。Closure Auditor 在新的审查上下文中从 Stage Contract 反向检查代码、测试、真实运行和遗留项；机器校验负责稳定 ID、任务覆盖、模板和证据路径。中断、恢复或上下文压缩后先加载 Program Charter、当前 Stage Contract 和当前任务 ID，再继续原任务。",
        "使用 Codex 长任务时，先用 Plan 明确 outcome、constraints 和 definition of done，再以 Goal 持续执行；AGENTS.md 保存仓库级不变量，Skill 保存可复用流程，SessionStart/PreCompact/Stop Hooks 负责恢复上下文和机械闭环检查。定时任务只做巡检或审计，不取代同一 Goal 的执行状态。",
    ]:
        anchor = insert_paragraph_after(anchor, text)

    annotation_table = table_with_headers(document, ["批注要点", "本稿处理"])
    annotation_rows = [
        ["Harness 概念混用；新自动流程不能直接替代旧流程", "区分工作流内 Harness、Evaluation Harness 和 Platform Harness；新 Profile 与旧路径并行实验，达到明确改进标准后才调整默认。"],
        ["验收自动化应可选，简单任务不应失去直接交付", "新增 Quick/Guided/Governed 三模式；Quick 允许需求直接生成并运行，Guided/Governed 按需要增加证据。"],
        ["积木点开不能只看到 JSON，应有真实配置界面", "新增 schema 驱动的领域配置器、专家 JSON、统一 WorkflowSpec 保存和典型积木配置要求。"],
        ["修改配置后应重验，但验收不应强制阻止发布", "行为修改使证据 stale 并建议重验；Quick/Guided 发布权属于用户，只有明确 Governed 策略执行硬门禁。"],
        ["失败用例要给建议并可一键修改", "失败生成结构化建议，并把节点、Trace、WorkflowSpec 和建议交回 Builder Team 的自然语言编辑预览。"],
        ["重型监管不能让当前工作流更难用", "监管按风险分层；优先恢复生成、编辑、运行体验，独立 Governance Console 服务工程和高风险场景。"],
        ["恢复 Lilies 的原始产品设定并处理自动演进偏移", "重写问题背景和最终路线；新增 Program Charter、稳定意图 ID、Stage Contract、偏移登记、Closure Auditor 与恢复协议。"],
        ["模板市场可以把好工作流封装为模块复用", "能力承载层新增可复用模块/子工作流与 Template/Module Registry，Builder 优先复用高阶模块。"],
        ["不是所有缺失能力都应做成积木", "新增五类承载位置和选择标准：积木、模块、Runtime/平台服务、平台硬控制、Connector/环境契约。"],
        ["层次模型问题可直接回答", "明确回答：对累积运行责任成立，对任意领域能力不成立；继续使用能力集合与运行闭包混合模型。"],
    ]
    while len(annotation_table.rows) > 1:
        annotation_table._tbl.remove(annotation_table.rows[-1]._tr)
    for row_values in annotation_rows:
        cells = annotation_table.add_row().cells
        for index, value in enumerate(row_values):
            replace_cell_text(cells[index], value)

    sources = table_with_headers(document, ["来源", "用于本报告的论点", "链接"])
    append_source_row(
        sources,
        "OpenAI - Long-running work",
        "长任务需要清晰 outcome、constraints、definition of done，并在同一 Goal 中持续执行和恢复。",
        "https://learn.chatgpt.com/docs/long-running-work",
    )
    append_source_row(
        sources,
        "OpenAI - AGENTS.md / Build skills / Hooks",
        "把仓库不变量、可复用流程、恢复和停止校验分别放入持久指令、Skill 与确定性生命周期脚本。",
        "https://learn.chatgpt.com/docs/agent-configuration/agents-md ; https://learn.chatgpt.com/docs/build-skills ; https://learn.chatgpt.com/docs/hooks",
    )

    document.core_properties.title = "Lilies 智能体需求能力模型与平台演进边界分析"
    document.core_properties.subject = "批注吸收与落地控制版"
    document.core_properties.comments = "Generated from the annotated v0.4.x revised report; inline annotations are integrated into the architecture baseline."
    validate_annotation_output(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
