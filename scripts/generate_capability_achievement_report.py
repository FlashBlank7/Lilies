from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/lilies_capability_boundary_achievement_review_2026-07-17.md"
OUTPUT = ROOT / "docs/lilies_capability_boundary_achievement_review_2026-07-17.docx"

BODY_FONT = "PingFang SC"
MONO_FONT = "SFMono-Regular"
INK = RGBColor(35, 43, 52)
NAVY = RGBColor(31, 78, 121)
TEAL = RGBColor(0, 112, 120)
MUTED = RGBColor(96, 108, 118)


def set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def add_inline(paragraph, text: str, *, size: float | None = None) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, BODY_FONT, size)
            run.font.color.rgb = INK
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, MONO_FONT, (size or 10.5) - 0.5)
            run.font.color.rgb = TEAL
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, BODY_FONT, size)
            run.bold = True
            run.font.color.rgb = INK
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, BODY_FONT, size)
        run.font.color.rgb = INK


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color in [
        ("Title", 24, NAVY),
        ("Heading 1", 17, NAVY),
        ("Heading 2", 13.5, TEAL),
        ("Heading 3", 11.5, NAVY),
    ]:
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "LILIES / CAPABILITY BOUNDARY ACHIEVEMENT REVIEW"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, MONO_FONT, 7.5)
        run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Lilies 能力边界成果审查  |  第 ")
    set_run_font(run, BODY_FONT, 8)
    run.font.color.rgb = MUTED
    add_page_field(footer)
    run = footer.add_run(" 页")
    set_run_font(run, BODY_FONT, 8)
    run.font.color.rgb = MUTED

    document.core_properties.title = "Lilies 能力边界设想落地成果审查报告"
    document.core_properties.subject = "v0.4.2-v0.4.10 implementation review"
    document.core_properties.author = "Lilies project"
    document.core_properties.keywords = "Lilies, capability boundary, v0.4.x, evaluation, governance"


def add_table(document: Document, raw_lines: list[str]) -> None:
    rows = parse_table([raw_lines[0], *raw_lines[2:]])
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        keep_row_together(table.rows[-1])
        if row_index == 0:
            repeat_header(table.rows[-1])
        for column_index in range(columns):
            value = values[column_index] if column_index < len(values) else ""
            cell = cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, "1F4E79")
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F2F6F8")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline(paragraph, value, size=8.2 if columns >= 5 else 8.8)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def build_document() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)

    index = 0
    first_heading = True
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            table_lines = [stripped, lines[index + 1].strip()]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_table(document, table_lines)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2)
            if first_heading:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_inline(paragraph, title, size=24)
                accent = document.add_paragraph()
                accent.paragraph_format.space_after = Pt(12)
                run = accent.add_run("REPORT / 2026-07-17 / v0.4.x")
                set_run_font(run, MONO_FONT, 9)
                run.font.color.rgb = TEAL
                first_heading = False
            else:
                paragraph = document.add_paragraph(style=f"Heading {level}")
                add_inline(paragraph, title)
            index += 1
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
            add_inline(paragraph, stripped[2:])
            index += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(3)
            add_inline(paragraph, numbered.group(2))
            index += 1
            continue

        paragraph = document.add_paragraph()
        if stripped.startswith("**审查") or stripped.startswith("**当前代码"):
            paragraph.paragraph_format.space_after = Pt(2)
        add_inline(paragraph, stripped)
        index += 1

    document.save(OUTPUT)

    reopened = Document(OUTPUT)
    fragments = [paragraph.text for paragraph in reopened.paragraphs]
    fragments.extend(
        cell.text
        for table in reopened.tables
        for row in table.rows
        for cell in row.cells
    )
    text = "\n".join(fragments)
    required = [
        "执行结论",
        "与能力边界报告路线逐项对照",
        "三个压力场景目前到底做到哪一步",
        "仍然不能对外承诺的事项",
        "741 passed, 85 xfailed, 0 failed",
    ]
    missing = [marker for marker in required if marker not in text]
    if missing:
        raise RuntimeError(f"generated report is missing required markers: {missing}")
    if len(reopened.tables) < 3:
        raise RuntimeError("generated report lost expected comparison tables")


if __name__ == "__main__":
    build_document()
